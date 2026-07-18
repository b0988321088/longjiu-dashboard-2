#!/usr/bin/env python3
"""Generate loan contract DOCX (Plan B, 300K, 5%, 12/31 final)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def fmt(run, font_name='標楷體', size=12, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_article_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('■ ' + title)
    fmt(run, size=12, bold=True)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)


def add_item(doc, text, style='List Number'):
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        fmt(run, size=12)


def add_indent(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        fmt(run, size=12)


# ===== PAGE 1 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('金錢借貸契約書')
fmt(run, size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('（方案B：按月還款、到期一筆清償暨本票擔保版）')
fmt(run, size=11)

doc.add_paragraph()

for party in ['貸與人（以下簡稱甲方）：__________________________',
              '借款人（以下簡稱乙方）：__________________________',
              '連帶保證人（以下簡稱丙方）：__________________________']:
    p = doc.add_paragraph()
    run = p.add_run(party)
    fmt(run, size=12)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('茲因乙方因業務週轉需要，向甲方借款，經雙方充分協議，同意條款如下，以資共同遵守：')
fmt(run, size=12)

doc.add_paragraph()
add_article_header(doc, '第一條：借貸明細與交付方式')

for item in [
    '一、借貸總金額：新台幣參拾萬元整（NT$ 300,000）。',
    '二、借貸期限：自民國 115 年 7 月 ____ 日起，至民國 115 年 12 月 31 日止。',
    '三、交付方式：甲方於本契約簽署後，將全額借款匯入乙方指定帳戶（合作金庫銀行東沙鹿分行，戶名：洛禾思健康發展有限公司）。匯款完成即視為交付完畢。',
]:
    add_item(doc, item)

doc.add_paragraph()
add_article_header(doc, '第二條：還款方式與利息計算（方案B）')

for item in [
    '一、利率：雙方約定本借款之年利率為 5%，利息按每月實際天數計算。',
    '二、每月還款：於借貸期限內，乙方每月償還新台幣 6,000 元，首期還款優先抵充利息，餘額抵充本金。每月還款日為每月 ____ 日。',
    '三、到期清償：乙方應於民國 115 年 12 月 31 日前，將剩餘未還本金及至該日為止之應計利息一次全額清償。',
    '四、還款帳號：',
]:
    add_item(doc, item)

add_indent(doc, '銀行：________　　分行：________　　帳號：________')

doc.add_paragraph()

# 還款明細表
p = doc.add_paragraph()
run = p.add_run('五、還款明細表（預計）')
fmt(run, size=12, bold=True)

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'

headers = ['月份', '每月繳款', '本金', '利息', '剩餘本金', '備註']
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = text
    for run in cell.paragraphs[0].runs:
        fmt(run, size=11, bold=True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

balance = 300000
for month in range(7, 13):
    if month == 12:
        interest = balance * 0.05 / 12
        principal = balance
        payment = balance + interest
        note = '到期一筆清'
    else:
        payment = 6000
        interest = balance * 0.05 / 12
        principal = 6000 - interest
        note = '按月繳'
    balance -= principal

    row = table.add_row().cells
    vals = [
        f'民國 {month} 月',
        f'{round(payment):,}',
        f'{round(principal):,}',
        f'{round(interest):,}',
        f'{round(balance):,}',
        note,
    ]
    for i, v in enumerate(vals):
        row[i].text = str(v)
        for run in row[i].paragraphs[0].runs:
            fmt(run, size=11)
        row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('★ 12 月 31 日到期需一筆清償尾款：277,201 元（本金 276,051 + 利息 1,150）')
fmt(run, size=12, bold=True)
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_paragraph()

add_article_header(doc, '第三條：擔保本票約定')

for item in [
    '一、為擔保乙方對甲方之本票債務，乙方及丙方應於簽約時共同簽發本票乙紙，交付甲方收執。',
    '二、本票記載如下：',
]:
    add_item(doc, item)

for item in [
    '票面金額：新台幣 300,000 元整',
    '發票人：洛禾思健康發展有限公司（蓋公司大小章）、陳宛琳（簽名蓋章）',
    '付款地：甲方指定址所',
    '到期日：民國 115 年 12 月 31 日',
]:
    add_indent(doc, item)

add_item(doc, '三、乙方及丙方如依約清償，甲方應無條件返還本票正本；如有違約，甲方得提示本票並依法請求付款。')

doc.add_page_break()

# ===== PAGE 2 =====
add_article_header(doc, '第四條：違約責任與加速條款（保護甲方條款）')

for item in [
    '一、失權效果（加速條款）：乙方如有任何一期未按時償還，即喪失期限利益，甲方得逕行主張全部債務即時到期，乙方及丙方應於甲方通知日起 5 日內一次清償全部剩餘本金、利息及一切費用。',
    '二、票據行使：乙方或丙方如有違約，甲方得隨時提示本票，並依法向法院聲請強制執行、拍賣或扣押乙方及丙方之一切財產，乙方及丙方絕無異議。',
    '三、逾期罰則：乙方逾期償還時，除仍按原約定利率計息外，並應按逾期未還金額之萬分之五，自逾期翌日起至清償日止計付違約金。',
]:
    add_item(doc, item)

doc.add_paragraph()
add_article_header(doc, '第五條：連帶保證責任')

add_item(doc, '一、丙方就乙方對甲方所負之本票債務（包括本金、利息、違約金、損害賠償及甲方因追償所生之律師費、訴訟費、執行費等）負連帶清償責任。')

add_item(doc, '二、甲方得分別或同時向丙方請求清償，無須先就乙方財產強制執行或選定執行順序；丙方亦不得主張先訴抗辯權。')

doc.add_paragraph()
add_article_header(doc, '第六條：管轄法院')

p = doc.add_paragraph('一、因本契約涉訟時，甲乙雙方合意以臺灣 ______ 地方法院為第一審管轄法院。')
for run in p.runs:
    fmt(run, size=12)

doc.add_paragraph()

# Special reminder box
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('█ 甲方簽約特別提醒')
fmt(run, size=12, bold=True)
run.font.color.rgb = RGBColor(0x8A, 0x5A, 0x00)

for r in [
    '1. 乙方簽約時，必須由「洛禾思健康發展有限公司」親蓋公司大小章，負責人陳宛琳親自簽名並蓋章。',
    '2. 丙方陳宛琳須親自簽名蓋章，並填寫身分證字號，以確保其個人負連帶清償責任。',
    '3. 本票發票人欄位需由「洛禾思健康發展有限公司」（大小章）與「陳宛琳」（簽名）共同簽發，確保本票債權對公司及個人均有效。',
]:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        fmt(run, size=11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('第 2 頁，共 3 頁')
fmt(run, size=10, color=(0x99, 0x99, 0x99))

doc.add_page_break()

# ===== PAGE 3 =====
p = doc.add_paragraph()
run = p.add_run('甲方（貸與人）：')
fmt(run, size=12, bold=True)

for f in ['姓名：____________', '身分證字號：____________', '地址：____________', '電話：____________']:
    p = doc.add_paragraph(f)
    for run in p.runs:
        fmt(run, size=12)

p = doc.add_paragraph('（簽名蓋章）')
for run in p.runs:
    fmt(run, size=11, color=(0x88, 0x88, 0x88))

p = doc.add_paragraph('────────────────────────────────────────────')
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
for run in p.runs:
    fmt(run, size=10, color=(0xCC, 0xCC, 0xCC))

p = doc.add_paragraph()
run = p.add_run('乙方（借款人）：')
fmt(run, size=12, bold=True)

for f in ['公司名稱：洛禾思健康發展有限公司', '統一編號：____________', '地址：____________', '負責人：陳宛琳']:
    p = doc.add_paragraph(f)
    for run in p.runs:
        fmt(run, size=12)

p = doc.add_paragraph('（公司大章與負責人小章）')
for run in p.runs:
    fmt(run, size=11, color=(0x88, 0x88, 0x88))

p = doc.add_paragraph('────────────────────────────────────────────')
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
for run in p.runs:
    fmt(run, size=10, color=(0xCC, 0xCC, 0xCC))

p = doc.add_paragraph()
run = p.add_run('丙方（連帶保證人）：')
fmt(run, size=12, bold=True)

for f in ['姓名：陳宛琳', '身分證字號：____________', '地址：____________', '電話：____________']:
    p = doc.add_paragraph(f)
    for run in p.runs:
        fmt(run, size=12)

p = doc.add_paragraph('（本人親自簽名蓋章）')
for run in p.runs:
    fmt(run, size=11, color=(0x88, 0x88, 0x88))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('中華民國 115 年 7 月 ____ 日')
fmt(run, size=12, bold=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('第 3 頁，共 3 頁')
fmt(run, size=10, color=(0x99, 0x99, 0x99))

out = Path('C:/Users/bot/Desktop/龍九系統/loan_contract_300k_b.docx')
doc.save(str(out))
print(f'✅ saved: {out}')
