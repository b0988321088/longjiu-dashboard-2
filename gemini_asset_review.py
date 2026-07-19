#!/usr/bin/env python3
"""Gemini review for asset_diff report."""
from __future__ import annotations
import json, os, re
from datetime import date
from pathlib import Path
from html.parser import HTMLParser

try:
    from dotenv import load_dotenv
    load_dotenv(Path.home() / "AppData" / "Local" / "hermes" / ".env")
except Exception:
    pass

import requests

BASE = Path(__file__).parent.resolve()
TODAY = Path(__file__).stem.split('_')[-1] if '_' in Path(__file__).stem else Path('.').cwd().name
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_BASE_URL = os.getenv("GEMINI_BASE_URL", "https://generativelanguage.googleapis.com/v1beta")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")

class TE(HTMLParser):
    def __init__(self):
        super().__init__()
        self.items = []
        self.skip = False
        self.depth = 0
    def handle_starttag(self, tag, attrs):
        if tag in {"style", "script", "head"}:
            self.skip = True
            self.depth += 1
    def handle_endtag(self, tag):
        if tag in {"style", "script", "head"}:
            self.depth -= 1
            if self.depth <= 0:
                self.skip = False
                self.depth = 0
    def handle_data(self, data):
        if not self.skip:
            self.items.append(data)

def to_text(path: Path) -> str:
    raw = path.read_text(encoding="utf-8")
    parser = TE()
    parser.feed(raw)
    txt = " ".join(parser.items)
    return re.sub(r"\s+", " ", txt).strip()[:4000]

def load_docx_text(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return " ".join(parts)[:4000]
    except Exception:
        return ""

def load_report() -> str:
    today = date.today().isoformat()
    html = BASE / f"asset_diff_{today}.html"
    docx = BASE / f"asset_diff_{today}.docx"
    if html.exists():
        return to_text(html)
    if docx.exists():
        return load_docx_text(docx)
    return ""

def review() -> dict:
    if not GEMINI_API_KEY:
        return {"status": "skipped", "reason": "GEMINI_API_KEY not set"}

    report = load_report()
    prompt = (
        "你是CIO-Gemini代理人。審查龍九資產變化對照報告。"
        "嚴格檢查：\n"
        "1. 數字計算正確性：覆蓋率必須是房租淨收+保守配息/月支出；cash必須用real_liquid_assets；保單質押=4M；國泰轉貸pending；房租淨收=80,100\n"
        "2. snapshot真值一致性：不硬編碼、不重複計算、金額未確認時只能寫'待確認'\n"
        "3. 保單明細正確性：僅顯示FL65，不得誤列FJ33為保單\n"
        "4. 30天歷史趨勢數據是否正確衔接\n"
        "5. 資產配置百分比計算是否有誤\n"
        "如有錯誤，務必指出錯誤點與正確值。\n\n"
        f"資產報告內容：\n{report}\n\n"
        '請用JSON回覆：{"status":"approved"或"rejected","issues":[{"point":"","description":"","fix":""}],"score":1-10,"summary":""}'
    )

    url = f"{GEMINI_BASE_URL}/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 2048, "responseMimeType": "application/json"},
    }
    try:
        r = requests.post(url, json=payload, timeout=30, headers={"Content-Type": "application/json"})
        if r.status_code != 200:
            return {"status": "error", "reason": f"HTTP {r.status_code}", "detail": r.text[:200]}
        candidate_texts = []
        for cand in r.json().get("candidates", []):
            content = cand.get("content", {})
            for part in content.get("parts", []):
                if "text" in part:
                    candidate_texts.append(part["text"])
        model_text = "\n".join(candidate_texts)
        try:
            result = json.loads(model_text)
            result["_raw"] = model_text[:500]
            return result
        except Exception:
            clean = re.sub(r"```(?:json)?\s*", "", model_text).strip()
            start = clean.find("{")
            if start >= 0:
                depth = 0
                end = -1
                for i, ch in enumerate(clean[start:], start):
                    if ch == "{":
                        depth += 1
                    elif ch == "}":
                        depth -= 1
                        if depth == 0:
                            end = i
                            break
                if end >= 0:
                    try:
                        result = json.loads(clean[start:end+1])
                        result["_raw"] = model_text[:500]
                        return result
                    except Exception:
                        pass
            raw_score = 0
            m = re.search(r'"score"\s*:\s*(\d+)', model_text)
            if m:
                raw_score = int(m.group(1))
            return {"status": "ok", "raw": model_text[:500], "score": raw_score}
    except Exception as exc:
        return {"status": "error", "reason": str(exc)}

if __name__ == "__main__":
    from datetime import date
    print(json.dumps(review(), ensure_ascii=False, indent=2))
